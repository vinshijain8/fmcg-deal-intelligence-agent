"""
exporter.py

Purpose:
---------
Export newsletter and data into
DOCX and JSON formats.
"""

import os
import json
import pandas as pd
from docx import Document


def export_files():

    os.makedirs("output", exist_ok=True)

    # -----------------------
    # Export JSON
    # -----------------------

    df = pd.read_csv("data/scored_news.csv")

    df.to_json(
        "output/scored_news.json",
        orient="records",
        indent=4
    )

    # -----------------------
    # Export Word
    # -----------------------

    with open(
        "output/newsletter.md",
        "r",
        encoding="utf-8"
    ) as f:

        newsletter = f.read()

    doc = Document()

    doc.add_heading(
        "FMCG Industry Intelligence Newsletter",
        level=1
    )

    doc.add_paragraph(newsletter)

    doc.save(
        "output/newsletter.docx"
    )

    print("Files exported successfully.")

    return True


if __name__ == "__main__":
    export_files()